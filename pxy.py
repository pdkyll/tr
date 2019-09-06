import os
import docx
from tr import *
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt
document=Document()
paragraph=document.add_paragraph('this is begining')
path = os.getcwd()
filepath=os.path.join(path,'imgs')
filelist=os.listdir(filepath)
filelist.sort()
#filelist=sorted(filelist,key=lambda i:int(i.split(".")[0]))
#print(filelist)
for filename in filelist:
    imagefile=os.path.join(filepath,filename)
    img_pil = Image.open(imagefile)
    MAX_SIZE = 2000
    if img_pil.height > MAX_SIZE or img_pil.width > MAX_SIZE:
        scale = max(img_pil.height / MAX_SIZE, img_pil.width / MAX_SIZE)
        new_width = int(img_pil.width / scale + 0.5)
        new_height = int(img_pil.height / scale + 0.5)
        img_pil = img_pil.resize((new_width, new_height), Image.BICUBIC)
    print(img_pil.width, img_pil.height)
#    img_pil    
    color_pil = img_pil.convert("RGB")
    gray_pil = img_pil.convert("L")
    rect_arr = detect(gray_pil, FLAG_RECT)
#    img_draw = ImageDraw.Draw(color_pil)
#    colors = [ 'red','green', 'blue', 'yellow', 'pink','black']
#    for i, rect in enumerate(rect_arr):
#        x, y, w, h = rect
#        img_draw.rectangle(
#            (x, y, x + w, y + h),
#            outline=colors[i % len(colors)],
#            )
#    color_pil
#    blank_pil = Image.new("L", img_pil.size, 255)
#    blank_draw = ImageDraw.Draw(blank_pil)
    results = run(gray_pil)    
    for line in results:
        x, y, w, h = line[0]
        txt = line[1]
#       font = ImageFont.truetype("msyh.ttf", max(int(h * 0.6), 14))
#        blank_draw.text(xy=(x, y), text=txt, font=font)
        paragraph.add_run(txt)

#blank_pil
print (filename)
document.save('test.docx')
print("ok")