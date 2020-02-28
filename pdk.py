import os
import docx
from tr import *
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE

document=Document()
paragraph=document.add_paragraph('-----this is begining-----')
path = os.getcwd()
filepath=os.path.join(path,'imgs')
filelist=os.listdir(filepath)
filelist.sort()
#filelist=sorted(filelist,key=lambda i:int(i.split(".")[0]))
style_index=0
filecount=0
print('filelist:',filelist)
for filename in filelist:
    filecount=filecount+1
    imagefile=os.path.join(filepath,filename)
    img_pil = Image.open(imagefile)
    MAX_SIZE = 2000
    if img_pil.height > MAX_SIZE or img_pil.width > MAX_SIZE:
        scale = max(img_pil.height / MAX_SIZE, img_pil.width / MAX_SIZE)
        new_width = int(img_pil.width / scale + 0.5)
        new_height = int(img_pil.height / scale + 0.5)
        img_pil = img_pil.resize((new_width, new_height), Image.BICUBIC)
    print(filecount,filename,img_pil.width, img_pil.height)
    #   color_pil = img_pil.convert("RGB")
    gray_pil = img_pil.convert("L")
    #    rect_arr = detect(gray_pil, FLAG_RECT)
    #    img_draw = ImageDraw.Draw(color_pil)
    #    colors = [ 'red','green', 'blue', 'yellow', 'pink','black']
    #    for i, rect in enumerate(rect_arr):
    #        x, y, w, h = rect
    #        img_draw.rectangle(
    #            (x, y, x + w, y + h),
    #            outline=colors[i % len(colors)],
    #            )
    #    color_pil
    #    blank_pil = Image.new("L", img_pil.size, 255)
    #    blank_draw = ImageDraw.Draw(blank_pil)
    results = run(gray_pil)   
    print('results')
    gray_pil.close()
    # color_pil.close()
    img_pil.close()#close imgfile !!!?
    # insert page sep
    # run1=paragraph.add_run()
    # run1.add_break(WD_BREAK.PAGE)
    for line in results:
        x, y, w, h = line[0]
        txt = line[1]
        print(txt)
        font = ImageFont.truetype("msyh.ttf", max(int(h * 0.6), 14))
        #blank_draw.text(xy=(x, y), text=txt, font=font)
        run2 = document.add_paragraph(txt)
        # style_index=style_index+1
        # style = document.styles.add_style('UserStyle%d' %style_index,WD_STYLE_TYPE.PARAGRAPH)
        # style.font.size=Pt(max(int(h * 0.6), 14))
        # run2.style=style
    
#blank_pil
print (filename)
document.save('output.docx')
print("ok")

